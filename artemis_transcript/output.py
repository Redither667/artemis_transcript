from pathlib import Path
from typing import override, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from artemis_transcript.core import OutputFuncSet


class MarkdownOutput(OutputFuncSet):
    def __init__(self, file: Path, *args, **kwargs):
        self.file = open(file, *args, **kwargs)

    def __del__(self):
        self.file.close()

    @override
    def write_story_line(self, text: str):
        self.file.write(f'# {text}\n')

    @override
    def write_save_title(self, text: str):
        self.file.write(f'## {text}\n')

    @override
    def write_select_title(self, text: str):
        self.file.write(f'### {text}\n')

    @override
    def write_text(self, name: Optional[str], text: str, face: Optional[str] = None):
        if name is None:
            assert face is None
            self.file.write(text)
        else:
            if face is None:
                self.file.write(f'{name}：{text}')
            else:
                self.file.write(f'{name}:')
                self.write_text_in_parenthesis(face)
                self.file.write(text)

    @override
    def write_text_in_parenthesis(self, text: str):
        self.file.write(f'（*{text}*）')

    @override
    def write_italic_text(self, text: str):
        self.file.write(f'*{text}*')

    @override
    def newline(self):
        self.file.write('  \n')

class DocxOutput(OutputFuncSet):
    def __init__(self, file: Path):
        self.path = file
        self.document = Document()
        self.paragraph = self.document.add_paragraph()
        style = self.document.styles['Normal']
        style.font.name = 'Times New Roman'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def __del__(self):
        self.document.save(str(self.path))

    @override
    def write_story_line(self, text: str):
        title = self.document.add_heading(text, level=0)
        title.alignment = 1
        self.paragraph = self.document.add_paragraph()

    @override
    def write_save_title(self, text: str):
        self.document.add_heading(text, level=1)
        self.paragraph = self.document.add_paragraph()

    @override
    def write_select_title(self, text: str):
        self.document.add_heading(text, level=2)
        self.paragraph = self.document.add_paragraph()

    @override
    def write_text(self, name: Optional[str], text: str, face: Optional[str] = None):
        if name is None:
            assert face is None
            self.paragraph.add_run(text)
        else:
            if face is None:
                self.paragraph.add_run(f'{name}：{text}')
            else:
                self.paragraph.add_run(f'{name}:')
                self.write_text_in_parenthesis(face)
                self.paragraph.add_run(text)

    @override
    def write_text_in_parenthesis(self, text: str):
        self.paragraph.add_run('（')
        self.write_italic_text(text)
        self.paragraph.add_run('）')

    @override
    def write_italic_text(self, text: str):
        run = self.paragraph.add_run(text)
        run.font.name = 'Times New Roman Italic'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

    @override
    def newline(self):
        self.paragraph.add_run('').add_break()